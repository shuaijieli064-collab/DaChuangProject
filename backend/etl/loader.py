import logging
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档对象"""
    content: str
    metadata: dict = field(default_factory=dict)
    source: str = ""  # 文件路径或来源标识


@dataclass
class Chunk:
    """文本切片对象"""
    content: str
    metadata: dict = field(default_factory=dict)
    chunk_id: str = ""
    embedding: list[float] | None = None


class DocumentLoader(ABC):
    """文档加载器基类"""

    @abstractmethod
    def load(self, source: str) -> list[Document]:
        """从源加载文档"""
        ...


class FileLoader(DocumentLoader):
    """单文件加载器"""

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def load(self, source: str) -> list[Document]:
        ext = os.path.splitext(source)[1].lower()
        loader = FILE_LOADERS.get(ext, TextLoader())
        return loader.load(source)


class TextLoader(DocumentLoader):
    """纯文本加载器"""

    def load(self, source: str, encoding: str = "utf-8") -> list[Document]:
        with open(source, encoding=encoding) as f:
            content = f.read()
        return [Document(content=content, metadata={"source": source}, source=source)]


class PDFLoader(DocumentLoader):
    """PDF 加载器"""

    def load(self, source: str) -> list[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader

        reader = PdfReader(source)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(Document(
                    content=text,
                    metadata={"source": source, "page": i + 1},
                    source=source,
                ))
        return pages


class DocxLoader(DocumentLoader):
    """Word 加载器"""

    def load(self, source: str) -> list[Document]:
        import docx
        doc = docx.Document(source)
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(content=content, metadata={"source": source}, source=source)]


# 文件扩展名 → 加载器映射
FILE_LOADERS: dict[str, DocumentLoader] = {
    ".txt": TextLoader(),
    ".md": TextLoader(),
    ".pdf": PDFLoader(),
    ".docx": DocxLoader(),
    ".doc": DocxLoader(),
}
